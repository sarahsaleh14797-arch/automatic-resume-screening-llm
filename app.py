import re
import sys
import json
import base64
import hashlib
import subprocess
import io
import html as html_lib
import shutil
from datetime import datetime
from pathlib import Path

import requests
import streamlit as st
import pandas as pd
from docx import Document

from src.llm.llm_client import generate_response
from src.ui.config import APP_NAME, LOGO_PATH, OLLAMA_BASE_URL, OLLAMA_MODEL, TAGLINE
from src.ui.theme import apply_theme


REPO_ROOT = Path(__file__).resolve().parent
DATA_DIR = REPO_ROOT / "data"
CVS_DIR = DATA_DIR / "samples" / "cvs"
JD_DIR = DATA_DIR / "samples" / "jd"
JD_HISTORY_DIR = JD_DIR / "history"
JD_FILE = JD_DIR / "job.txt"

OUT_DIR = DATA_DIR / "outputs"
EXTRACTED_DIR = OUT_DIR / "extracted_text"
CHUNKS_DIR = OUT_DIR / "chunks"
RANKING_DIR = OUT_DIR / "ranking"

VECTORSTORE_DIR = DATA_DIR / "vectorstore"


def ensure_dirs():
    CVS_DIR.mkdir(parents=True, exist_ok=True)
    JD_DIR.mkdir(parents=True, exist_ok=True)
    JD_HISTORY_DIR.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    EXTRACTED_DIR.mkdir(parents=True, exist_ok=True)
    CHUNKS_DIR.mkdir(parents=True, exist_ok=True)
    RANKING_DIR.mkdir(parents=True, exist_ok=True)


def sanitize_filename(name: str) -> str:
    name = Path(name).name
    name = re.sub(r"[^\w.\- ()\[\]]+", "_", name, flags=re.UNICODE).strip()
    return name or "uploaded_file"


def stable_key(s: str) -> str:
    return hashlib.md5(s.encode("utf-8", errors="ignore")).hexdigest()


def save_uploaded_cvs(uploaded_files):
    saved = []
    for uf in uploaded_files:
        fname = sanitize_filename(uf.name)
        dest = CVS_DIR / fname
        if dest.exists():
            stem, suf = dest.stem, dest.suffix
            i = 1
            while True:
                cand = CVS_DIR / f"{stem}_{i}{suf}"
                if not cand.exists():
                    dest = cand
                    break
                i += 1
        dest.write_bytes(uf.getbuffer())
        saved.append(dest.name)
    return saved


def clear_sample_cvs():
    for p in CVS_DIR.glob("*"):
        if p.is_file():
            p.unlink(missing_ok=True)


def write_jd(text: str):
    JD_FILE.write_text(text.strip() + "\n", encoding="utf-8")


def save_jd_version(text: str):
    JD_HISTORY_DIR.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = JD_HISTORY_DIR / f"job_{ts}.txt"
    path.write_text(text.strip() + "\n", encoding="utf-8")
    return path


def _try_unlink(path: Path, errors: list[str]):
    try:
        if path.exists():
            path.unlink()
    except Exception as e:
        errors.append(f"{path}: {type(e).__name__}: {e}")


def _clear_dir_files(dir_path: Path, pattern: str, errors: list[str]):
    if not dir_path.exists():
        return
    for p in dir_path.glob(pattern):
        if p.is_file():
            _try_unlink(p, errors)


def clear_outputs_fresh_run():
    errors = []

    _clear_dir_files(EXTRACTED_DIR, "*.txt", errors)
    _clear_dir_files(CHUNKS_DIR, "*.txt", errors)
    _clear_dir_files(RANKING_DIR, "*.csv", errors)
    _clear_dir_files(RANKING_DIR, "*.json", errors)

    try:
        if VECTORSTORE_DIR.exists():
            shutil.rmtree(VECTORSTORE_DIR)
    except Exception as e:
        errors.append(f"{VECTORSTORE_DIR}: {type(e).__name__}: {e}")

    for k in ["last_run_logs", "last_run_time"]:
        if k in st.session_state:
            del st.session_state[k]

    return errors


def run_cmd(cmd_list, cwd: Path):
    proc = subprocess.Popen(
        cmd_list,
        cwd=str(cwd),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    while True:
        line = proc.stdout.readline()
        if not line:
            break
        yield line
    proc.wait()
    if proc.returncode != 0:
        raise RuntimeError(f"Command failed ({proc.returncode}): {' '.join(cmd_list)}")


def load_json_any(path: Path):
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8", errors="replace"))
    except Exception:
        return None


def inject_logo_as_data_url():
    p = Path(LOGO_PATH)
    if not p.exists():
        return None
    data = p.read_bytes()
    b64 = base64.b64encode(data).decode("utf-8")
    ext = p.suffix.lower().replace(".", "")
    mime = "png" if ext == "png" else "svg+xml" if ext == "svg" else "png"
    return f"data:image/{mime};base64,{b64}"


def find_latest(pattern: str, base: Path) -> Path | None:
    files = list(base.rglob(pattern))
    if not files:
        return None
    files.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return files[0]


def normalize_explanations(data):
    if data is None:
        return []
    if isinstance(data, list):
        return [x for x in data if isinstance(x, dict)]
    if isinstance(data, dict):
        out = []
        for k, v in data.items():
            if isinstance(v, dict):
                vv = dict(v)
                vv.setdefault("cv_source", k)
                out.append(vv)
        return out
    return []


def split_bullets(text: str):
    lines = [l.strip() for l in (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    items = []
    for l in lines:
        if not l:
            continue
        l = re.sub(r"^\s*[-*•]\s*", "", l)
        l = re.sub(r"^\s*\d+[\).\s-]+\s*", "", l)
        if l:
            items.append(l)
    return items


def parse_llm_analysis(text: str):
    t = (text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not t:
        return None

    m = re.search(
        r"(?ms)^\s*1\)\s*(.*?)^\s*2\)\s*(.*?)^\s*3\)\s*(.*?)^\s*4\)\s*(.*?)^\s*5\)\s*(.*)\s*$",
        t,
    )
    if m:
        summary = m.group(1).strip()
        strengths = m.group(2).strip()
        gaps = m.group(3).strip()
        rec = m.group(4).strip()
        just = m.group(5).strip()
        return {
            "summary": summary,
            "strengths": split_bullets(strengths),
            "gaps": split_bullets(gaps),
            "recommendation": rec,
            "justification": just,
            "raw": t,
        }

    return {"summary": "", "strengths": [], "gaps": [], "recommendation": "", "justification": "", "raw": t}


def esc(s: str) -> str:
    return html_lib.escape(str(s or ""))


def make_candidate_docx_bytes(item: dict, parsed: dict | None) -> bytes:
    doc = Document()
    title = item.get("cv_source") or item.get("cv_name") or "candidate"
    doc.add_heading(f"CV Evaluation - {title}", level=1)

    doc.add_paragraph(f"Score: {item.get('score', '')}")
    doc.add_paragraph(f"Status: {item.get('status', '')}")
    if item.get("error"):
        doc.add_paragraph(f"Error: {item.get('error')}")

    if parsed:
        if parsed.get("recommendation"):
            doc.add_heading("Recommendation", level=2)
            doc.add_paragraph(parsed["recommendation"])

        if parsed.get("summary"):
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(parsed["summary"])

        strengths = parsed.get("strengths") or []
        if strengths:
            doc.add_heading("Strengths", level=2)
            for x in strengths:
                doc.add_paragraph(str(x), style="List Bullet")

        gaps = parsed.get("gaps") or []
        if gaps:
            doc.add_heading("Gaps / Missing skills", level=2)
            for x in gaps:
                doc.add_paragraph(str(x), style="List Bullet")

        if parsed.get("justification"):
            doc.add_heading("Justification", level=2)
            doc.add_paragraph(parsed["justification"])
    else:
        raw = item.get("llm_analysis", "") or ""
        if raw.strip():
            doc.add_heading("Raw LLM output", level=2)
            doc.add_paragraph(raw)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_candidate_html(item: dict, parsed: dict | None) -> str:
    title = item.get("cv_source") or item.get("cv_name") or "candidate"
    score = item.get("score", "")
    status = item.get("status", "")
    took = item.get("took_sec", "")
    err = item.get("error", "")

    rec = parsed.get("recommendation", "") if parsed else ""
    summ = parsed.get("summary", "") if parsed else ""
    strengths = parsed.get("strengths", []) if parsed else []
    gaps = parsed.get("gaps", []) if parsed else []
    just = parsed.get("justification", "") if parsed else ""
    raw = parsed.get("raw", "") if parsed else (item.get("llm_analysis", "") or "")

    strengths_html = "".join([f"<li>{esc(x)}</li>" for x in strengths])
    gaps_html = "".join([f"<li>{esc(x)}</li>" for x in gaps])

    return f"""<!doctype html>
<html>
<head>
<meta charset="utf-8"/>
<title>{esc(title)} - CV Evaluation</title>
<style>
body{{font-family:Arial, sans-serif; margin:24px; color:#111}}
h1{{margin:0 0 8px 0}}
.small{{color:#444; font-size:13px; margin-bottom:16px}}
.section{{margin-top:18px}}
pre{{white-space:pre-wrap; background:#f4f4f4; padding:12px; border-radius:8px}}
ul{{margin-top:6px}}
</style>
</head>
<body>
<h1>{esc(title)}</h1>
<div class="small">Score: {esc(score)} | Status: {esc(status)} | Took: {esc(took)} | Error: {esc(err)}</div>

<div class="section"><h2>Recommendation</h2><div>{esc(rec)}</div></div>
<div class="section"><h2>Summary</h2><div>{esc(summ)}</div></div>
<div class="section"><h2>Strengths</h2><ul>{strengths_html}</ul></div>
<div class="section"><h2>Gaps / Missing skills</h2><ul>{gaps_html}</ul></div>
<div class="section"><h2>Justification</h2><div>{esc(just)}</div></div>
<div class="section"><h2>Raw LLM output</h2><pre>{esc(raw)}</pre></div>

</body>
</html>"""


def render_llm_item(item: dict):
    score = item.get("score", None)
    status = item.get("status", None)
    took = item.get("took_sec", None)
    err = item.get("error", None)
    analysis = item.get("llm_analysis", "") or item.get("llm_analysis_text", "") or item.get("analysis", "")

    meta_cols = st.columns(4)
    meta_cols[0].write(f"score: {score}" if score is not None else "score: -")
    meta_cols[1].write(f"status: {status}" if status else "status: -")
    meta_cols[2].write(f"took_sec: {took}" if took is not None else "took_sec: -")
    meta_cols[3].write(f"error: {err}" if err else "error: -")

    parsed = parse_llm_analysis(analysis) if (analysis or "").strip() else None

    title = item.get("cv_source") or item.get("cv_name") or "candidate"
    safe_name = re.sub(r"[^\w.\- ()\[\]]+", "_", str(title)).strip() or "candidate"

    docx_bytes = make_candidate_docx_bytes(item, parsed)
    html_text = make_candidate_html(item, parsed)

    c1, c2 = st.columns(2)
    with c1:
        st.download_button(
            "Download DOCX",
            data=docx_bytes,
            file_name=f"{safe_name}_evaluation.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    with c2:
        st.download_button(
            "Download HTML",
            data=html_text.encode("utf-8"),
            file_name=f"{safe_name}_evaluation.html",
            mime="text/html",
            use_container_width=True,
        )

    if status in ["timeout", "error"] and not (analysis or "").strip():
        return

    if not parsed:
        st.code((analysis or "").strip())
        return

    rec = (parsed.get("recommendation") or "").strip()
    if rec:
        st.write("Recommendation")
        st.write(rec)

    summ = (parsed.get("summary") or "").strip()
    if summ:
        st.write("Summary")
        st.write(summ)

    strengths = parsed.get("strengths") or []
    if strengths:
        st.write("Strengths")
        st.markdown("\n".join([f"- {x}" for x in strengths]))

    gaps = parsed.get("gaps") or []
    if gaps:
        st.write("Gaps / Missing skills")
        st.markdown("\n".join([f"- {x}" for x in gaps]))

    just = (parsed.get("justification") or "").strip()
    if just:
        st.write("Justification")
        st.write(just)

    raw_key = f"raw_{stable_key(str(item.get('cv_source','')) + str(item.get('cv_name','')))}"
    if st.checkbox("Show raw LLM output", key=raw_key, value=False):
        st.code(parsed.get("raw", ""))


st.set_page_config(page_title=APP_NAME, layout="wide")
apply_theme()
ensure_dirs()

logo_data_url = inject_logo_as_data_url()

if logo_data_url is None:
    st.error(f"Logo not found at: {LOGO_PATH}")
else:
    st.markdown(
        f"""
        <style>
          .sira-header {{
            display: flex;
            align-items: center;
            gap: 18px;
            padding: 22px 8px 10px 8px;
            margin-top: 6px;
          }}
          .sira-logo {{
            width: 220px;
            height: auto;
            display: block;
          }}
          .sira-title {{
            font-size: 42px;
            line-height: 1.05;
            font-weight: 800;
            color: var(--sira-gold);
            margin: 0;
            padding: 0;
          }}
          .sira-tagline {{
            font-size: 18px;
            line-height: 1.35;
            font-weight: 600;
            color: var(--sira-gold);
            margin-top: 6px;
          }}
          .sira-subline {{
            font-size: 15px;
            line-height: 1.35;
            font-weight: 500;
            color: #ffffff;
            margin-top: 6px;
            opacity: 0.95;
          }}
          @media (max-width: 900px) {{
            .sira-header {{ flex-direction: column; align-items: flex-start; }}
            .sira-logo {{ width: 190px; }}
            .sira-title {{ font-size: 34px; }}
          }}
        </style>
        <div class="sira-header">
          <div class="sira-left"><img class="sira-logo" src="{logo_data_url}" /></div>
          <div class="sira-right">
            <div class="sira-title">{APP_NAME}</div>
            <div class="sira-tagline">{TAGLINE}</div>
            <div class="sira-subline">Local LLM via Ollama • RAG Pipeline • Ranking + Explainability</div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

st.divider()
tabs = st.tabs(["Upload", "Run", "Results", "System Check"])

with tabs[0]:
    st.subheader("Upload CVs and Job Description")

    uploaded_cvs = st.file_uploader("Upload CV files (PDF/DOCX)", type=["pdf", "docx"], accept_multiple_files=True)
    existing_jd = JD_FILE.read_text(encoding="utf-8", errors="replace") if JD_FILE.exists() else ""
    jd_text = st.text_area("Paste Job Description", height=220, value=existing_jd)

    c1, c2, c3, c4 = st.columns([1, 1, 1, 1])
    with c1:
        save_btn = st.button("Save Inputs", use_container_width=True)
    with c2:
        save_ver_btn = st.button("Save JD as Version", use_container_width=True)
    with c3:
        clear_cvs_btn = st.button("Clear CVs Folder", use_container_width=True)
    with c4:
        clear_outputs_btn = st.button("Clear Outputs (Fresh Run)", use_container_width=True)

    if save_ver_btn:
        content = jd_text.strip()
        if not content:
            st.error("Job Description is empty")
        else:
            p = save_jd_version(content)
            st.success(f"Saved JD version: {p}")

    if clear_cvs_btn:
        clear_sample_cvs()
        st.success("CVs folder cleared")
        st.rerun()

    if clear_outputs_btn:
        errs = clear_outputs_fresh_run()
        if errs:
            st.error("Could not clear some files. Close any opened CSV/JSON (Excel) then try again.")
            st.code("\n".join(errs))
        else:
            st.success("Outputs cleared. Next run will be a fresh run.")
            st.rerun()

    if save_btn:
        if not jd_text.strip():
            st.error("Job Description is empty")
        else:
            write_jd(jd_text)
            saved = save_uploaded_cvs(uploaded_cvs) if uploaded_cvs else []
            st.success("Inputs saved")
            st.write("Saved JD:", str(JD_FILE))
            if saved:
                st.write("Saved CVs:")
                for s in saved:
                    st.write(f"- {s}")
            else:
                st.warning("No CVs uploaded in this save")

    st.divider()
    st.subheader("Saved CVs")
    files = sorted([p for p in CVS_DIR.glob("*") if p.is_file()], key=lambda x: x.name.lower())
    if not files:
        st.info("No CV files saved")
    else:
        for p in files:
            a, b = st.columns([5, 1])
            with a:
                st.write(p.name)
            with b:
                if st.button("Delete", key=f"del_{stable_key(p.name)}"):
                    p.unlink(missing_ok=True)
                    st.rerun()

with tabs[1]:
    st.subheader("Run Screening")

    run_btn = st.button("Run Pipeline", type="primary")

    jd_ok = JD_FILE.exists() and JD_FILE.read_text(encoding="utf-8", errors="replace").strip() != ""
    cv_ok = any(CVS_DIR.glob("*.pdf")) or any(CVS_DIR.glob("*.docx"))

    st.write("Inputs Status")
    st.write(f"- JD file: {'OK' if jd_ok else 'Missing/Empty'}")
    st.write(f"- CVs folder: {'OK' if cv_ok else 'No CVs found'}")

    if run_btn:
        if not jd_ok:
            st.error("Missing/empty JD. Save JD first.")
            st.stop()
        if not cv_ok:
            st.error("No CVs found. Upload and save CVs first.")
            st.stop()

        steps = [
            [sys.executable, "src/01_ingest/extract_text.py"],
            [sys.executable, "src/02_preprocessing/chunk_text.py"],
            [sys.executable, "src/04_vectorstore/reset_vectorstore.py"],
            [sys.executable, "src/03_embeddings/embed_and_store.py"],
            [sys.executable, "src/07_ranking/rank_cvs.py"],
            [sys.executable, "-m", "src.llm.explain_with_llm"],
        ]

        log_box = st.empty()
        logs = ""
        prog = st.progress(0)

        try:
            for i, cmd in enumerate(steps, start=1):
                st.write(f"Step {i}/{len(steps)}: {' '.join(cmd)}")
                for line in run_cmd(cmd, cwd=REPO_ROOT):
                    logs += line
                    log_box.code(logs[-5000:])
                prog.progress(int(i / len(steps) * 100))
            st.success("Pipeline finished")
        except Exception as e:
            st.error(f"Run failed: {e}")
        finally:
            prog.empty()

        st.session_state["last_run_logs"] = logs
        st.session_state["last_run_time"] = str(pd.Timestamp.now())

with tabs[2]:
    st.subheader("Results")

    latest_csv = find_latest("ranking_results*.csv", RANKING_DIR)
    if latest_csv and latest_csv.exists():
        try:
            df = pd.read_csv(latest_csv)
            st.write(str(latest_csv))
            st.dataframe(df, use_container_width=True)
        except Exception as e:
            st.error(f"Failed to read ranking CSV: {e}")
    else:
        st.info("No ranking CSV found")

    st.divider()

    llm_path = RANKING_DIR / "llm_explanations.json"
    llm_data = load_json_any(llm_path)
    items = normalize_explanations(llm_data)

    if not items:
        st.info("No llm_explanations.json found or file is empty")
    else:
        items.sort(
            key=lambda x: float(x.get("score", 0.0)) if str(x.get("score", "")).strip() != "" else -1.0,
            reverse=True,
        )
        for item in items:
            title = item.get("cv_source") or item.get("cv_name") or "candidate"
            with st.expander(str(title), expanded=False):
                render_llm_item(item)

    if st.checkbox("Show last run logs", value=False) and "last_run_logs" in st.session_state:
        st.divider()
        st.code(st.session_state["last_run_logs"][-8000:])

with tabs[3]:
    st.subheader("System Check")

    ollama_ok = False
    model_ok = False

    try:
        r = requests.get(f"{OLLAMA_BASE_URL.rstrip('/')}/api/tags", timeout=3)
        if r.status_code == 200:
            ollama_ok = True
            data = r.json()
            models = [m.get("name", "") for m in data.get("models", [])]
            model_ok = any(OLLAMA_MODEL == m or m.startswith(OLLAMA_MODEL) for m in models)
    except Exception:
        pass

    st.write("Runtime Status")
    st.write(f"- Ollama service: {'OK' if ollama_ok else 'Not reachable'}")
    st.write(f"- Required model ({OLLAMA_MODEL}): {'Available' if model_ok else 'Missing'}")

    if not ollama_ok:
        st.error("Ollama is not running. Start Ollama and refresh this page.")
    elif not model_ok:
        st.warning(f"Model not found. Run: ollama pull {OLLAMA_MODEL}")
    else:
        st.success("System prerequisites look good.")

    test_prompt = st.text_area("Enter a test prompt:", height=120)

    if st.button("Run LLM Test"):
        if test_prompt.strip() == "":
            st.warning("Please enter a prompt.")
        else:
            with st.spinner("Generating response..."):
                result = generate_response(test_prompt)
            st.write(result)